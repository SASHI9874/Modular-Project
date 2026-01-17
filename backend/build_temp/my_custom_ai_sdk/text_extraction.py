import os
import csv
import fitz
from docx import Document
import openpyxl

class TextExtractor:
    def __init__(self):
        pass

    def _clean_markdown(self, text: str) -> str:
        """Removes excessive newlines but preserves paragraph structure."""
        if not text: return ""
        # Remove multiple empty lines greater than 2
        import re
        return re.sub(r'\n{3,}', '\n\n', text.strip())

    def _extract_pdf_markdown(self, file_path: str) -> str:
        """
        Extracts PDF text. Uses font size heuristic to detect Headers.
        """
        md_lines = []
        try:
            with fitz.open(file_path) as doc:
                for page in doc:
                    blocks = page.get_text("dict")["blocks"]
                    for block in blocks:
                        if "lines" in block:
                            for line in block["lines"]:
                                for span in line["spans"]:
                                    text = span["text"].strip()
                                    if not text: continue
                                    
                                    # Heuristic: If font size > 15, treat as Header
                                    if span["size"] > 18:
                                        md_lines.append(f"\n# {text}\n")
                                    elif span["size"] > 14:
                                        md_lines.append(f"\n## {text}\n")
                                    elif span["flags"] & 2 ** 4: # Bold flag in PyMuPDF
                                        md_lines.append(f"**{text}** ")
                                    else:
                                        md_lines.append(f"{text} ")
                            md_lines.append("\n\n") # Paragraph break after block
            return "".join(md_lines)
        except Exception as e:
            return f"> Error reading PDF: {str(e)}"

    def _extract_docx_markdown(self, file_path: str) -> str:
        """
        Extracts DOCX preserving Headers, Lists, and Tables.
        """
        try:
            doc = Document(file_path)
            md_output = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text: continue

                style_name = para.style.name.lower()
                
                # Map Styles to Markdown
                if 'heading 1' in style_name:
                    md_output.append(f"\n# {text}\n")
                elif 'heading 2' in style_name:
                    md_output.append(f"\n## {text}\n")
                elif 'heading 3' in style_name:
                    md_output.append(f"\n### {text}\n")
                elif 'list' in style_name:
                    md_output.append(f"* {text}")
                elif 'quote' in style_name:
                    md_output.append(f"> {text}")
                else:
                    md_output.append(text)
                
                md_output.append("\n\n")

            # Basic Table Extraction for DOCX
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    cells = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                    rows.append(f"| {' | '.join(cells)} |")
                
                if rows:
                    # Add separator line
                    header_len = len(table.rows[0].cells)
                    sep = f"| {' | '.join(['---'] * header_len)} |"
                    rows.insert(1, sep)
                    md_output.append("\n" + "\n".join(rows) + "\n")

            return "".join(md_output)
        except Exception as e:
            return f"> Error reading DOCX: {str(e)}"

    def _extract_spreadsheet_markdown(self, file_path: str, is_csv=False) -> str:
        """
        Converts CSV/Excel rows into Markdown Tables.
        """
        rows = []
        try:
            if is_csv:
                with open(file_path, 'r', encoding='utf-8') as f:
                    reader = csv.reader(f)
                    rows = list(reader)
            else:
                # Excel
                wb = openpyxl.load_workbook(file_path, data_only=True)
                ws = wb.active
                for row in ws.iter_rows(values_only=True):
                    # Convert None to empty string
                    rows.append([str(cell) if cell is not None else "" for cell in row])
            
            if not rows: return ""

            # Format as Markdown Table
            md_lines = []
            
            # Header
            headers = rows[0]
            md_lines.append(f"| {' | '.join(headers)} |")
            
            # Separator
            md_lines.append(f"| {' | '.join(['---'] * len(headers))} |")
            
            # Body
            for row in rows[1:]:
                # Ensure row length matches header
                padded_row = row + [""] * (len(headers) - len(row))
                md_lines.append(f"| {' | '.join(padded_row)} |")

            return "\n".join(md_lines)

        except Exception as e:
            return f"> Error reading Spreadsheet: {str(e)}"

    def run(self, file_path: str) -> str:
        """
        Main entry point. Returns text in Markdown format.
        """
        if not os.path.exists(file_path):
            return "> Error: File not found."

        ext = file_path.lower().split('.')[-1]
        
        result = ""
        
        if ext == 'pdf':
            result = self._extract_pdf_markdown(file_path)
        elif ext == 'docx':
            result = self._extract_docx_markdown(file_path)
        elif ext in ['xls', 'xlsx']:
            result = self._extract_spreadsheet_markdown(file_path, is_csv=False)
        elif ext == 'csv':
            result = self._extract_spreadsheet_markdown(file_path, is_csv=True)
        elif ext == 'txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                result = f.read()
        else:
            # Fallback for unsupported types
            return f"> Warning: Format '.{ext}' is not supported for rich extraction."

        return self._clean_markdown(result)